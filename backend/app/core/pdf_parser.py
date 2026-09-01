import pymupdf


def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file using pymupdf."""
    text = ""
    try:
        doc = pymupdf.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {str(e)}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX resume (paragraphs plus table cells).

    The frontend accepts .docx uploads alongside PDFs, so the analysis
    pipeline has to handle both (FR-1)."""
    import docx  # python-docx

    try:
        document = docx.Document(file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {str(e)}")

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_resume_text(file_path: str) -> str:
    """Dispatches to the PDF or DOCX extractor based on the file extension."""
    if file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    return extract_text_from_pdf(file_path)
